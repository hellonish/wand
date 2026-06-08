"""Tests for profile document ingestion."""

from io import BytesIO

import pytest

from engine.profile import (
    FileType,
    LinkKind,
    ProfileDocumentInput,
    ProfileIngestor,
    SourceType,
    TextBlock,
    ingest_document,
    ingest_documents,
)


def _input(filename: str, data: bytes, content_type: str = "text/plain", declared=None):
    """Build a document input for ingestion tests."""

    return ProfileDocumentInput(
        filename=filename,
        content_type=content_type,
        file_bytes=data,
        declared_source_type=declared,
    )


def test_text_ingestion_extracts_blocks_and_links_without_semantic_type():
    """Extract text blocks and written links from text input."""

    document = ingest_document(
        _input(
            "resume.txt",
            b"Experience\nData Analyst at Acme\nSkills\nPython\nEducation\nBS\nhttps://linkedin.com/in/jane",
        )
    )

    assert document.file_type == FileType.TXT
    assert document.source_type == SourceType.OTHER
    assert document.text_blocks
    assert document.links[0].kind == LinkKind.OTHER
    assert document.links[0].link_id
    assert document.links[0].source.value == "text"
    assert document.links[0].block_id


def test_html_ingestion_gets_embedded_and_written_links():
    """Extract both href links and written links from HTML input."""

    html = b"""
    <html>
      <head><title>Jane Portfolio</title></head>
      <body>
        <h1>Portfolio</h1>
        <h2>Projects</h2>
        <a href="https://github.com/jane/project">Forecasting repo</a>
        <p>Live demo: jane-demo.vercel.app</p>
      </body>
    </html>
    """

    document = ingest_document(_input("portfolio.html", html, "text/html"))

    urls = {link.url for link in document.links}
    assert document.source_type == SourceType.OTHER
    assert "https://github.com/jane/project" in urls
    assert "https://jane-demo.vercel.app" in urls


def test_pdf_ingestion_preserves_page_numbers():
    """Preserve PDF page numbers on blocks and link evidence."""

    fitz = pytest.importorskip("fitz")
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Jane Doe Resume\nExperience\nData Analyst\nhttps://example.com")
    data = pdf.tobytes()
    pdf.close()

    document = ingest_document(_input("resume.pdf", data, "application/pdf"))

    assert document.file_type == FileType.PDF
    assert document.metadata.page_count == 1
    assert document.text_blocks[0].page_number == 1
    assert document.links[0].page_number == 1


def test_docx_ingestion_extracts_paragraphs_tables_when_dependency_exists():
    """Extract paragraphs, tables, and links from DOCX input."""

    docx = pytest.importorskip("docx")

    document = docx.Document()
    document.add_heading("Projects", level=1)
    document.add_paragraph("Built a forecasting dashboard.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Demo"
    table.cell(0, 1).text = "https://demo.example.com"
    buffer = BytesIO()
    document.save(buffer)

    ingested = ingest_document(
        _input(
            "projects.docx",
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )

    assert ingested.file_type == FileType.DOCX
    assert ingested.source_type == SourceType.OTHER
    assert any("Demo | https://demo.example.com" in block.text for block in ingested.text_blocks)
    assert any(link.url == "https://demo.example.com" for link in ingested.links)


def test_docx_ingestion_does_not_duplicate_paragraph_hyperlink_relationships():
    """Keep DOCX paragraph hyperlinks from being recaptured by the relationship sweep."""

    docx = pytest.importorskip("docx")
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    document = docx.Document()
    paragraph = document.add_paragraph("Portfolio: ")
    rel_id = paragraph.part.relate_to(
        "https://portfolio.example.com",
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Portfolio site"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    buffer = BytesIO()
    document.save(buffer)

    ingested = ingest_document(
        _input(
            "portfolio.docx",
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )

    matching_links = [link for link in ingested.links if link.url == "https://portfolio.example.com"]
    assert len(matching_links) == 1
    assert matching_links[0].block_id


def test_link_block_matching_returns_none_when_evidence_does_not_overlap():
    """Avoid assigning the last block when a captured link has no matching evidence."""

    ingestor = ProfileIngestor()
    blocks = [TextBlock(block_id="doc:block:0", text="Experience at Acme")]

    assert ingestor._match_block_for_link("unrelated link context", "external", blocks) is None


def test_doc_file_is_rejected():
    """Reject unsupported .doc files."""

    with pytest.raises(ValueError, match="Unsupported .doc"):
        ingest_document(_input("resume.doc", b"old binary doc"))


def test_ingest_documents_marks_duplicate_sha():
    """Mark duplicate file content during batch ingestion."""

    documents = ingest_documents(
        [
            _input("a.txt", b"Experience\nSkills\nEducation\nhttps://example.com"),
            _input("b.txt", b"Experience\nSkills\nEducation\nhttps://example.com"),
        ]
    )

    assert documents[0].metadata.sha256 == documents[1].metadata.sha256
    assert documents[1].metadata.duplicate_of == documents[0].document_id


def test_declared_source_type_wins():
    """Prefer a declared source type over inferred classification."""

    document = ingest_document(_input("file.txt", b"Experience\nSkills\nEducation", declared=SourceType.COVER_LETTER))

    assert document.source_type == SourceType.COVER_LETTER
