"""Straightforward profile document ingestion."""

import hashlib
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import fitz
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn

from .models import (
    CapturedLink,
    DocumentMetadata,
    FileType,
    IngestedProfileDocument,
    LinkKind,
    LinkSource,
    ProfileDocumentInput,
    SourceType,
    TextBlock,
)


URL_RE = re.compile(
    r"\b(?:(?:https?://|www\.)[^\s<>()\[\]{}\"']+|"
    r"(?:[A-Za-z0-9-]+\.)+(?:com|org|net|io|dev|app|me|co|ai|edu|gov)"
    r"(?:/[^\s<>()\[\]{}\"']*)?)",
    re.IGNORECASE,
)
EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")

EXTENSION_TO_TYPE = {
    ".pdf": FileType.PDF,
    ".html": FileType.HTML,
    ".htm": FileType.HTML,
    ".docx": FileType.DOCX,
    ".txt": FileType.TXT,
}


class ProfileIngestor:
    """Read supported files into document blocks and links."""

    def ingest_document(self, document_input: ProfileDocumentInput) -> IngestedProfileDocument:
        """Ingest one document without character-level provenance."""

        extension = Path(document_input.filename).suffix.lower()
        if extension == ".doc":
            raise ValueError("Unsupported .doc file. Upload .docx, PDF, HTML, or TXT.")

        sha256 = hashlib.sha256(document_input.file_bytes).hexdigest()
        document_id = sha256[:16]
        file_type = self._file_type(document_input.filename, document_input.content_type)
        blocks, embedded_links, metadata_values, warnings = self._read(document_id, file_type, document_input.file_bytes)
        links = self._assign_link_ids(document_id, self._dedupe_links([*embedded_links, *self._links_from_blocks(blocks)]))
        metadata = DocumentMetadata(
            filename=document_input.filename,
            content_type=document_input.content_type,
            extension=extension,
            size_bytes=len(document_input.file_bytes),
            sha256=sha256,
            block_count=len(blocks),
            link_count=len(links),
            **metadata_values,
        )
        source_type = document_input.declared_source_type or SourceType.OTHER

        return IngestedProfileDocument(
            document_id=document_id,
            source_type=source_type,
            file_type=file_type,
            metadata=metadata,
            text_blocks=blocks,
            links=links,
            warnings=warnings,
        )

    def ingest_documents(self, document_inputs: List[ProfileDocumentInput]) -> List[IngestedProfileDocument]:
        """Ingest many documents and mark duplicate file content."""

        documents: List[IngestedProfileDocument] = []
        seen: Dict[str, str] = {}
        for document_input in document_inputs:
            document = self.ingest_document(document_input)
            if document.metadata.sha256 in seen:
                document.metadata.duplicate_of = seen[document.metadata.sha256]
                document.warnings.append(f"Duplicate file content detected: {seen[document.metadata.sha256]}")
            else:
                seen[document.metadata.sha256] = document.document_id
            documents.append(document)
        return documents

    def _read(
        self,
        document_id: str,
        file_type: FileType,
        file_bytes: bytes,
    ) -> Tuple[List[TextBlock], List[CapturedLink], Dict[str, int | str | None], List[str]]:
        """Dispatch bytes to the reader for the detected file type."""

        if file_type == FileType.PDF:
            return self._read_pdf(document_id, file_bytes)
        if file_type == FileType.HTML:
            return self._read_html(document_id, file_bytes)
        if file_type == FileType.DOCX:
            return self._read_docx(document_id, file_bytes)
        return self._read_text(document_id, file_bytes)

    def _read_pdf(
        self,
        document_id: str,
        file_bytes: bytes,
    ) -> Tuple[List[TextBlock], List[CapturedLink], Dict[str, int | str | None], List[str]]:
        """Extract text and embedded links from PDF pages."""

        blocks: List[TextBlock] = []
        links: List[CapturedLink] = []
        warnings: List[str] = []
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            for page_number, page in enumerate(pdf, start=1):
                text = page.get_text("text").strip()
                page_block_start = len(blocks)
                if text:
                    self._append_text_blocks(document_id, blocks, text, page_number)
                else:
                    warnings.append(f"Page {page_number} has no extractable text")
                page_blocks = blocks[page_block_start:]
                for link in page.get_links():
                    url = link.get("uri")
                    if url:
                        label, context, block = self._pdf_link_context(page, link, page_blocks)
                        links.append(
                            self._link(
                                url,
                                label or url,
                                block.block_id if block else None,
                                page_number,
                                context,
                                source=LinkSource.PDF_EMBEDDED,
                                heading_path=block.heading_path if block else [],
                            )
                        )
            return blocks, links, {"page_count": pdf.page_count}, warnings

    def _read_html(
        self,
        document_id: str,
        file_bytes: bytes,
    ) -> Tuple[List[TextBlock], List[CapturedLink], Dict[str, int | str | None], List[str]]:
        """Extract visible HTML text and anchor links."""

        soup = BeautifulSoup(self._decode(file_bytes), "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()

        blocks: List[TextBlock] = []
        links: List[CapturedLink] = []
        heading_path: List[str] = []
        for element in soup.find_all(["h1", "h2", "h3", "p", "li", "td", "th", "a", "div"]):
            if element.name == "div" and not self._is_kept_div(element):
                continue
            text = self._element_text(element)
            href = element.get("href") if element.name == "a" else None
            if element.name in {"h1", "h2", "h3"} and text:
                heading_path = heading_path[: int(element.name[1]) - 1] + [text]
            block_id = None
            if text:
                block = self._block(document_id, len(blocks), text, None, heading_path)
                blocks.append(block)
                block_id = block.block_id
            if href:
                context = self._html_link_context(element)
                nearest_block = self._match_block_for_link(context or text or href, text or href, blocks)
                if nearest_block:
                    block_id = nearest_block.block_id
                    link_heading_path = nearest_block.heading_path
                else:
                    link_heading_path = heading_path
                links.append(
                    self._link(
                        href,
                        text or href,
                        block_id,
                        None,
                        context or text or None,
                        source=LinkSource.HTML_HREF,
                        heading_path=link_heading_path,
                    )
                )

        title = soup.title.string.strip() if soup.title and soup.title.string else None
        return blocks, links, {"title": title}, []

    def _read_docx(
        self,
        document_id: str,
        file_bytes: bytes,
    ) -> Tuple[List[TextBlock], List[CapturedLink], Dict[str, int | str | None], List[str]]:
        """Extract paragraphs and table rows from DOCX."""

        doc = Document(BytesIO(file_bytes))
        blocks: List[TextBlock] = []
        links: List[CapturedLink] = []
        captured_relationship_ids: set[str] = set()
        heading_path: List[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name.lower() if paragraph.style else ""
            if style.startswith("heading"):
                heading_path = [text]
            block = self._block(document_id, len(blocks), text, None, heading_path)
            blocks.append(block)
            links.extend(self._docx_paragraph_links(doc, paragraph, block, captured_relationship_ids))

        for table in doc.tables:
            for row in table.rows:
                text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if text:
                    blocks.append(self._block(document_id, len(blocks), text, None, heading_path))

        for rel_id, rel in doc.part.rels.items():
            if rel_id in captured_relationship_ids:
                continue
            if getattr(rel, "is_external", False):
                target = getattr(rel, "target_ref", "")
                if target:
                    links.append(self._link(target, target, None, None, source=LinkSource.DOCX_RELATIONSHIP))

        return blocks, links, {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}, []

    def _read_text(
        self,
        document_id: str,
        file_bytes: bytes,
    ) -> Tuple[List[TextBlock], List[CapturedLink], Dict[str, int | str | None], List[str]]:
        """Extract text from plain text."""

        blocks: List[TextBlock] = []
        text = self._decode(file_bytes).strip()
        if text:
            blocks.append(self._block(document_id, len(blocks), self._clean_text(text), None))
        return blocks, [], {}, []

    def _append_text_blocks(
        self,
        document_id: str,
        blocks: List[TextBlock],
        text: str,
        page_number: Optional[int],
    ) -> None:
        """Append one raw text block for source evidence."""

        clean = self._clean_text(text)
        if clean:
            blocks.append(self._block(document_id, len(blocks), clean, page_number))

    def _block(
        self,
        document_id: str,
        index: int,
        text: str,
        page_number: Optional[int],
        heading_path: Optional[List[str]] = None,
    ) -> TextBlock:
        """Create a text block without semantic classification."""

        return TextBlock(
            block_id=f"{document_id}:block:{index}",
            text=text,
            page_number=page_number,
            heading_path=heading_path or [],
        )

    def _links_from_blocks(self, blocks: Iterable[TextBlock]) -> List[CapturedLink]:
        """Extract written URLs and emails from text blocks."""

        links: List[CapturedLink] = []
        for block in blocks:
            for match in URL_RE.finditer(block.text):
                if match.start() > 0 and block.text[match.start() - 1] == "@":
                    continue
                raw = match.group(0)
                url = self._normalize_url(raw)
                links.append(
                    self._link(
                        url,
                        raw,
                        block.block_id,
                        block.page_number,
                        self._context(block.text, match.start(), match.end()),
                        source=LinkSource.TEXT,
                        heading_path=block.heading_path,
                    )
                )
            for match in EMAIL_RE.finditer(block.text):
                raw = match.group(0)
                links.append(
                    self._link(
                        f"mailto:{raw}",
                        raw,
                        block.block_id,
                        block.page_number,
                        self._context(block.text, match.start(), match.end()),
                        source=LinkSource.TEXT,
                        heading_path=block.heading_path,
                    )
                )
        return links

    def _link(
        self,
        url: str,
        label: str,
        block_id: Optional[str],
        page_number: Optional[int],
        context: Optional[str] = None,
        source: LinkSource = LinkSource.TEXT,
        heading_path: Optional[List[str]] = None,
    ) -> CapturedLink:
        """Create a normalized captured link record."""

        normalized = self._normalize_url(url)
        return CapturedLink(
            url=normalized,
            kind=LinkKind.OTHER,
            label=label,
            context=context,
            block_id=block_id,
            page_number=page_number,
            source=source,
            heading_path=heading_path or [],
        )

    def _pdf_link_context(
        self,
        page,
        link: Dict,
        page_blocks: List[TextBlock],
    ) -> Tuple[Optional[str], Optional[str], Optional[TextBlock]]:
        """Extract label, nearby context, and nearest block for a PDF link annotation."""

        rect_value = link.get("from")
        if not rect_value:
            return None, None, None
        rect = fitz.Rect(rect_value)
        label = self._clean_text(page.get_textbox(rect))
        if not label:
            label = self._words_in_rect(page, rect)
        expanded = fitz.Rect(rect.x0 - 120, rect.y0 - 28, rect.x1 + 220, rect.y1 + 28)
        context = self._clean_text(page.get_textbox(expanded)) or label
        block = self._match_block_for_link(context or "", label or "", page_blocks)
        if not block:
            return label, context, None
        if block and (not context or len(context) < len(label or "")):
            context = self._context_from_block(block.text, label)
        return label, context, block

    def _words_in_rect(self, page, rect) -> Optional[str]:
        """Return words intersecting a PDF rectangle."""

        words = []
        for word in page.get_text("words"):
            word_rect = fitz.Rect(word[:4])
            if word_rect.intersects(rect):
                words.append(str(word[4]))
        return self._clean_text(" ".join(words)) or None

    def _html_link_context(self, element) -> Optional[str]:
        """Return useful surrounding text for an HTML anchor."""

        parent = element.find_parent(
            lambda tag: tag.name in {"article", "section", "li", "p"}
            or bool(set(tag.get("class", [])) & {"paper", "project", "project-card", "card"})
        )
        context_element = parent or element.parent or element
        text = self._clean_text(context_element.get_text("\n", strip=True))
        return text[:1800] if text else None

    def _docx_paragraph_links(
        self,
        doc: Document,
        paragraph,
        block: TextBlock,
        captured_relationship_ids: set[str],
    ) -> List[CapturedLink]:
        """Extract external hyperlink relationships from one DOCX paragraph."""

        links: List[CapturedLink] = []
        for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
            rel_id = hyperlink.get(qn("r:id"))
            if not rel_id or rel_id not in doc.part.rels:
                continue
            rel = doc.part.rels[rel_id]
            if not getattr(rel, "is_external", False):
                continue
            label = self._clean_text(" ".join(node.text for node in hyperlink.xpath(".//w:t") if node.text))
            url = getattr(rel, "target_ref", "")
            if url:
                captured_relationship_ids.add(rel_id)
                links.append(
                    self._link(
                        url,
                        label or url,
                        block.block_id,
                        block.page_number,
                        block.text,
                        source=LinkSource.DOCX_RELATIONSHIP,
                        heading_path=block.heading_path,
                    )
                )
        return links

    def _match_block_for_link(
        self,
        context: str,
        label: str,
        blocks: List[TextBlock],
    ) -> Optional[TextBlock]:
        """Find the source block closest to a captured link."""

        if not blocks:
            return None
        context_tokens = self._token_set(context)
        label_tokens = self._token_set(label)
        best: Optional[TextBlock] = None
        best_score = 0
        for block in reversed(blocks):
            block_tokens = self._token_set(block.text)
            score = len(block_tokens & context_tokens) * 3 + len(block_tokens & label_tokens)
            if score > best_score:
                best = block
                best_score = score
        return best if best_score else None

    def _context_from_block(self, text: str, label: Optional[str]) -> str:
        """Use the source block as context, bounded around the label when present."""

        if label and label in text:
            start = text.index(label)
            return self._context(text, start, start + len(label), 220)
        return text[:600]

    def _file_type(self, filename: str, content_type: str) -> FileType:
        """Resolve supported file type from extension or MIME type."""

        extension = Path(filename).suffix.lower()
        if extension in EXTENSION_TO_TYPE:
            return EXTENSION_TO_TYPE[extension]
        content_type = content_type.lower()
        if "pdf" in content_type:
            return FileType.PDF
        if "html" in content_type:
            return FileType.HTML
        if "wordprocessingml.document" in content_type:
            return FileType.DOCX
        if content_type.startswith("text/"):
            return FileType.TXT
        raise ValueError(f"Unsupported file extension: {extension or '<none>'}")

    def _normalize_url(self, url: str) -> str:
        """Normalize written URLs without rewriting anchors or local documents."""

        clean = url.rstrip(".,;:)")
        if clean.startswith(("mailto:", "#", "/", "./", "../")):
            return clean
        if clean.lower().endswith((".pdf", ".docx", ".doc", ".txt")) and not clean.startswith(("http://", "https://")):
            return clean
        if clean.startswith("www."):
            return f"https://{clean}"
        if not clean.startswith(("http://", "https://")) and "." in clean:
            return f"https://{clean}"
        return clean

    def _dedupe_links(self, links: List[CapturedLink]) -> List[CapturedLink]:
        """Remove duplicate link occurrences within one ingested document."""

        seen = set()
        result = []
        for link in links:
            key = (
                link.kind.value,
                link.url.lower().rstrip("/"),
                (link.label or "").lower().strip(),
                link.block_id,
                (link.context or "").lower().strip(),
            )
            if key in seen:
                continue
            seen.add(key)
            result.append(link)
        return result

    def _assign_link_ids(self, document_id: str, links: List[CapturedLink]) -> List[CapturedLink]:
        """Assign stable link ids after deduplication."""

        return [link.model_copy(update={"link_id": f"{document_id}:link:{index}"}) for index, link in enumerate(links)]

    def _context(self, text: str, start: int, end: int, window: int = 160) -> str:
        """Return a bounded nearby text snippet."""

        return text[max(0, start - window) : min(len(text), end + window)].strip()

    def _element_text(self, element) -> str:
        """Extract visible element text with useful separators for rich containers."""

        separator = "\n" if element.name == "div" else " "
        return self._clean_text(element.get_text(separator, strip=True))

    def _is_kept_div(self, element) -> bool:
        """Keep only divs that contain meaningful portfolio card text."""

        classes = set(element.get("class", []))
        return bool(classes & {"paper", "paper-title", "paper-authors", "paper-venue", "paper-links"})

    def _clean_text(self, text: Optional[str]) -> str:
        """Normalize whitespace while preserving line boundaries."""

        if not text:
            return ""
        lines = [" ".join(line.strip().split()) for line in text.splitlines() if line.strip()]
        return "\n".join(lines)

    def _token_set(self, text: str) -> set[str]:
        """Return normalized tokens useful for fuzzy block matching."""

        return {token for token in re.findall(r"[a-z0-9][a-z0-9+-]{2,}", text.lower()) if token not in {"https", "http", "www", "com", "org", "net"}}

    def _decode(self, file_bytes: bytes) -> str:
        """Decode bytes using common text encodings."""

        for encoding in ("utf-8", "utf-16"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("latin-1")


def ingest_document(document_input: ProfileDocumentInput) -> IngestedProfileDocument:
    """Ingest one profile document using the default ingestor."""

    return ProfileIngestor().ingest_document(document_input)


def ingest_documents(document_inputs: List[ProfileDocumentInput]) -> List[IngestedProfileDocument]:
    """Ingest many profile documents using the default ingestor."""

    return ProfileIngestor().ingest_documents(document_inputs)
